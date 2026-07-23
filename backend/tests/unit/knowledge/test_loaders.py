import io
from unittest.mock import patch

import docx
import pytest

from src.core.exceptions import ValidationException
from src.knowledge.loaders.docx_loader import DocxLoader
from src.knowledge.loaders.pdf_loader import PDFLoader
from src.knowledge.loaders.registry import get_loader, resolve_content_type
from src.knowledge.loaders.text_loader import TextLoader


def test_text_loader_decodes_bytes():
    loader = TextLoader()

    assert loader.load(b"hello world") == "hello world"


def test_docx_loader_round_trips_paragraphs():
    document = docx.Document()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("Second paragraph.")

    buffer = io.BytesIO()
    document.save(buffer)

    loader = DocxLoader()
    text = loader.load(buffer.getvalue())

    assert "First paragraph." in text
    assert "Second paragraph." in text


def test_pdf_loader_joins_page_text():
    fake_page_1 = type("Page", (), {"extract_text": lambda self: "Page one"})()
    fake_page_2 = type("Page", (), {"extract_text": lambda self: "Page two"})()
    fake_reader = type("Reader", (), {"pages": [fake_page_1, fake_page_2]})()

    with patch(
        "src.knowledge.loaders.pdf_loader.PdfReader",
        return_value=fake_reader,
    ):
        text = PDFLoader().load(b"fake pdf bytes")

    assert text == "Page one\n\nPage two"


def test_get_loader_raises_on_unsupported_type():
    with pytest.raises(ValidationException):
        get_loader("image/png")


def test_resolve_content_type_falls_back_to_extension():
    assert resolve_content_type("notes.md", None) == "text/markdown"
    # An unrecognized provided content type is overridden by the extension.
    assert resolve_content_type("report.pdf", "application/octet-stream") == (
        "application/pdf"
    )
    # A recognized provided content type is trusted as-is.
    assert resolve_content_type("notes.txt", "text/plain") == "text/plain"
    # No extension and no usable content type falls back to octet-stream.
    assert resolve_content_type("noextension", None) == "application/octet-stream"
